import uuid
import io
import os
from datetime import datetime, timezone
from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from typing import Optional
from services.report_service import ReportService
from middleware.auth_middleware import get_current_user
from utils.logger import get_logger
def _to_dict(obj) -> dict:
    try:
        return obj.model_dump()
    except AttributeError:
        return obj.dict()

router  = APIRouter(prefix="/reports", tags=["Reports"])
service = ReportService()
logger  = get_logger(__name__)


class ReportRequest(BaseModel):
    report_type:      str = "custom"
    format:           str = "pdf"
    date_from:        Optional[str] = None
    date_to:          Optional[str] = None
    severity:         Optional[str] = None
    status:           Optional[str] = None
    region:           Optional[str] = None
    username:         Optional[str] = None
    source_ip:        Optional[str] = None
    min_risk_score:   Optional[int] = 0
    asset_type:       Optional[str] = None


@router.post("/generate")
def generate_report(body: ReportRequest, current_user=Depends(get_current_user)):
    """
    Generate a security report from live AWS data.
    Returns the report in the requested format.
    """
    filters = {
        "date_from":      body.date_from,
        "date_to":        body.date_to,
        "severity":       body.severity,
        "status":         body.status,
        "region":         body.region,
        "username":       body.username,
        "source_ip":      body.source_ip,
        "min_risk_score": body.min_risk_score,
        "asset_type":     body.asset_type,
        "report_type":    body.report_type,
    }
    filters = {k: v for k, v in filters.items() if v is not None and v != 0}

    try:
        data      = service.gather_report_data(filters)
        report_id = f"RPT-{uuid.uuid4().hex[:8].upper()}"
        fmt       = body.format.lower()
        now_str   = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
        filename  = f"soc_report_{body.report_type}_{now_str}"

        if fmt == "txt":
            content  = _generate_txt(data, body.report_type, current_user.username)
            service.save_report_record(report_id, filename, current_user.username, "txt", body.report_type, filters)
            return StreamingResponse(
                io.BytesIO(content.encode("utf-8")),
                media_type="text/plain",
                headers={"Content-Disposition": f"attachment; filename={filename}.txt"}
            )

        elif fmt == "md":
            content  = _generate_markdown(data, body.report_type, current_user.username)
            service.save_report_record(report_id, filename, current_user.username, "md", body.report_type, filters)
            return StreamingResponse(
                io.BytesIO(content.encode("utf-8")),
                media_type="text/markdown",
                headers={"Content-Disposition": f"attachment; filename={filename}.md"}
            )

        elif fmt == "pdf":
            content = _generate_pdf(data, body.report_type, current_user.username)
            service.save_report_record(report_id, filename, current_user.username, "pdf", body.report_type, filters)
            return StreamingResponse(
                io.BytesIO(content),
                media_type="application/pdf",
                headers={"Content-Disposition": f"attachment; filename={filename}.pdf"}
            )

        elif fmt == "docx":
            content = _generate_docx(data, body.report_type, current_user.username)
            service.save_report_record(report_id, filename, current_user.username, "docx", body.report_type, filters)
            return StreamingResponse(
                io.BytesIO(content),
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={filename}.docx"}
            )

        else:
            raise HTTPException(400, f"Unsupported format: {fmt}. Use: pdf, docx, txt, md")

    except Exception as e:
        logger.error(f"Report generation failed: {e}")
        raise HTTPException(500, f"Report generation failed: {str(e)}")


@router.get("/preview")
def preview_report(
    report_type:   str = Query("custom"),
    date_from:     Optional[str] = Query(None),
    date_to:       Optional[str] = Query(None),
    severity:      Optional[str] = Query(None),
    status:        Optional[str] = Query(None),
    region:        Optional[str] = Query(None),
    username:      Optional[str] = Query(None),
    min_risk_score: int = Query(0),
    current_user=Depends(get_current_user)
):
    """Preview report data as JSON — used by frontend to show report content."""
    filters = {k: v for k, v in {
        "date_from": date_from, "date_to": date_to,
        "severity": severity, "status": status,
        "region": region, "username": username,
        "min_risk_score": min_risk_score,
        "report_type": report_type,
    }.items() if v is not None and v != 0}
    data = service.gather_report_data(filters)
    # Return preview — truncate large lists
    data["threats"]    = data["threats"][:20]
    data["incidents"]  = data["incidents"][:20]
    data["audit_logs"] = data["audit_logs"][:20]
    data["assets"]     = data["assets"][:20]
    return data


@router.get("/history")
def get_report_history(current_user=Depends(get_current_user)):
    return service.get_report_history()


@router.delete("/{report_id}")
def delete_report(report_id: str, current_user=Depends(get_current_user)):
    ok = service.delete_report(report_id)
    if not ok:
        raise HTTPException(500, "Failed to delete report")
    return {"message": f"Report {report_id} deleted"}


# ═══════════════════════════════════════════════════════════
# Report Generators
# ═══════════════════════════════════════════════════════════

def _generate_txt(data: dict, report_type: str, generated_by: str) -> str:
    ai   = data.get("ai_summary", {})
    anal = data.get("analytics", {})
    sev  = anal.get("severity_breakdown", {})
    lines = [
        "=" * 70,
        "   AWS CLOUD SECURITY OPERATIONS CENTER",
        f"   {report_type.upper().replace('_',' ')} REPORT",
        "=" * 70,
        f"Generated By : {generated_by}",
        f"Generated At : {data.get('generated_at','')}",
        f"Security Score: {ai.get('security_score','N/A')}/100",
        f"Risk Level   : {ai.get('risk_level','N/A')}",
        "",
        "─" * 70,
        "EXECUTIVE SUMMARY",
        "─" * 70,
        ai.get("executive_summary", "N/A"),
        "",
        "─" * 70,
        "THREAT SUMMARY",
        "─" * 70,
        f"Total Threats  : {anal.get('total_threats', 0)}",
        f"Critical       : {sev.get('CRITICAL', 0)}",
        f"High           : {sev.get('HIGH', 0)}",
        f"Medium         : {sev.get('MEDIUM', 0)}",
        f"Low            : {sev.get('LOW', 0)}",
        f"Avg Risk Score : {anal.get('average_risk_score', 0)}/100",
        "",
        "─" * 70,
        "INCIDENT SUMMARY",
        "─" * 70,
        f"Total Incidents  : {anal.get('total_incidents', 0)}",
        f"Open             : {anal.get('open_incidents', 0)}",
        f"Resolved         : {anal.get('resolved_incidents', 0)}",
        "",
        "─" * 70,
        "TOP THREAT ACTORS",
        "─" * 70,
    ]
    for u in anal.get("top_users", []):
        lines.append(f"  {u['name']:30s} {u['count']} events")

    lines += ["", "─" * 70, "TOP SOURCE IPs", "─" * 70]
    for ip in anal.get("top_source_ips", []):
        lines.append(f"  {ip['name']:30s} {ip['count']} events")

    lines += ["", "─" * 70, "CRITICAL FINDINGS", "─" * 70]
    for cf in ai.get("critical_findings", [])[:5]:
        lines.append(f"  • {cf}")

    lines += ["", "─" * 70, "RECOMMENDATIONS", "─" * 70]
    for i, rec in enumerate(ai.get("recommendations", []), 1):
        lines.append(f"  {i}. {rec}")

    lines += ["", "─" * 70, "RECENT THREATS (Top 20)", "─" * 70]
    for t in data.get("threats", [])[:20]:
        lines.append(
            f"  [{t.get('severity','?'):8s}] {t.get('event_type','?'):40s} "
            f"| {t.get('username','?'):15s} | {t.get('source_ip','?'):15s} | {t.get('timestamp','?')[:19]}"
        )

    lines += [
        "", "=" * 70,
        "AWS Cloud SOC Platform — Confidential Security Report",
        "=" * 70,
    ]
    return "\n".join(lines)


def _generate_markdown(data: dict, report_type: str, generated_by: str) -> str:
    ai   = data.get("ai_summary", {})
    anal = data.get("analytics", {})
    sev  = anal.get("severity_breakdown", {})
    now  = data.get("generated_at", "")

    lines = [
        f"# AWS Cloud SOC — {report_type.upper().replace('_',' ')} REPORT",
        "",
        f"| Field | Value |",
        f"|---|---|",
        f"| Generated By | {generated_by} |",
        f"| Generated At | {now} |",
        f"| Security Score | **{ai.get('security_score','N/A')}/100** |",
        f"| Risk Level | **{ai.get('risk_level','N/A')}** |",
        "",
        "---",
        "",
        "## Executive Summary",
        "",
        ai.get("executive_summary", "N/A"),
        "",
        "---",
        "",
        "## Threat Summary",
        "",
        f"| Severity | Count |",
        f"|---|---|",
        f"| 🔴 Critical | {sev.get('CRITICAL', 0)} |",
        f"| 🟠 High | {sev.get('HIGH', 0)} |",
        f"| 🟡 Medium | {sev.get('MEDIUM', 0)} |",
        f"| 🔵 Low | {sev.get('LOW', 0)} |",
        f"| **Total** | **{anal.get('total_threats', 0)}** |",
        "",
        f"> Average Risk Score: **{anal.get('average_risk_score', 0)}/100**",
        "",
        "---",
        "",
        "## Incident Summary",
        "",
        f"| Status | Count |",
        f"|---|---|",
        f"| Open | {anal.get('open_incidents', 0)} |",
        f"| Resolved | {anal.get('resolved_incidents', 0)} |",
        f"| **Total** | **{anal.get('total_incidents', 0)}** |",
        "",
        "---",
        "",
        "## Critical Findings",
        "",
    ]
    for cf in ai.get("critical_findings", [])[:5]:
        lines.append(f"- ⚠️ {cf}")

    lines += ["", "---", "", "## Top Threat Actors", ""]
    for u in anal.get("top_users", []):
        lines.append(f"- `{u['name']}` — {u['count']} events")

    lines += ["", "## Top Source IPs", ""]
    for ip in anal.get("top_source_ips", []):
        lines.append(f"- `{ip['name']}` — {ip['count']} events")

    lines += ["", "## Top AWS Services", ""]
    for svc in anal.get("top_services", []):
        lines.append(f"- `{svc['name']}` — {svc['count']} events")

    lines += ["", "---", "", "## Recommendations", ""]
    for i, rec in enumerate(ai.get("recommendations", []), 1):
        lines.append(f"{i}. {rec}")

    lines += ["", "---", "", "## Recent Threats", "",
              "| Threat ID | Event | Severity | User | Source IP | Time |",
              "|---|---|---|---|---|---|"]
    for t in data.get("threats", [])[:20]:
        lines.append(
            f"| {t.get('threat_id','')} | {t.get('event_type','')} "
            f"| {t.get('severity','')} | {t.get('username','')} "
            f"| {t.get('source_ip','')} | {t.get('timestamp','')[:19]} |"
        )

    lines += [
        "", "---", "",
        "*AWS Cloud SOC Platform — Confidential Security Report*",
        f"*{now}*",
    ]
    return "\n".join(lines)


def _generate_pdf(data: dict, report_type: str, generated_by: str) -> bytes:
    """Generate PDF using reportlab."""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.colors import HexColor, white, black
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer, Table,
            TableStyle, HRFlowable, PageBreak
        )
        from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

        buffer = io.BytesIO()
        doc    = SimpleDocTemplate(buffer, pagesize=A4,
                                   topMargin=2*cm, bottomMargin=2*cm,
                                   leftMargin=2*cm, rightMargin=2*cm)

        styles = getSampleStyleSheet()
        ai     = data.get("ai_summary", {})
        anal   = data.get("analytics", {})
        sev    = anal.get("severity_breakdown", {})

        # Colors
        dark_bg   = HexColor("#0d1117")
        cyan      = HexColor("#4cc9f0")
        red       = HexColor("#ff4d6d")
        orange    = HexColor("#ff8c42")
        yellow    = HexColor("#ffd166")
        green     = HexColor("#06d6a0")
        light_txt = HexColor("#e2e8f0")
        gray      = HexColor("#6b7280")

        h1 = ParagraphStyle("h1", parent=styles["Heading1"], textColor=cyan, fontSize=18, spaceAfter=6)
        h2 = ParagraphStyle("h2", parent=styles["Heading2"], textColor=cyan, fontSize=13, spaceAfter=4)
        body = ParagraphStyle("body", parent=styles["Normal"], textColor=black, fontSize=9, spaceAfter=4, leading=14)
        small = ParagraphStyle("small", parent=styles["Normal"], textColor=gray, fontSize=8, spaceAfter=2)
        center = ParagraphStyle("center", parent=styles["Normal"], alignment=TA_CENTER, textColor=gray, fontSize=8)

        story = []

        # ── Cover ──────────────────────────────────────────
        story.append(Spacer(1, 2*cm))
        story.append(Paragraph("🛡 AWS CLOUD SECURITY OPERATIONS CENTER", h1))
        story.append(Paragraph(f"{report_type.upper().replace('_',' ')} REPORT", h1))
        story.append(HRFlowable(width="100%", color=cyan, thickness=1))
        story.append(Spacer(1, 0.5*cm))

        meta = [
            ["Generated By", generated_by],
            ["Generated At", data.get("generated_at","")],
            ["Security Score", f"{ai.get('security_score','N/A')}/100"],
            ["Risk Level",    ai.get("risk_level","N/A")],
            ["Total Threats", str(anal.get("total_threats",0))],
            ["Open Incidents",str(anal.get("open_incidents",0))],
        ]
        meta_table = Table(meta, colWidths=[5*cm, 11*cm])
        meta_table.setStyle(TableStyle([
            ("FONTSIZE",    (0,0),(-1,-1), 9),
            ("TEXTCOLOR",   (0,0),(0,-1),  cyan),
            ("TEXTCOLOR",   (1,0),(1,-1),  black),
            ("FONTNAME",    (0,0),(0,-1),  "Helvetica-Bold"),
            ("ROWBACKGROUNDS",(0,0),(-1,-1),["#f8fafc","#f1f5f9"]),
            ("GRID",        (0,0),(-1,-1), 0.5, HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 5),
        ]))
        story.append(meta_table)
        story.append(Spacer(1, 0.5*cm))
        story.append(HRFlowable(width="100%", color=HexColor("#e2e8f0"), thickness=0.5))

        # ── Executive Summary ──────────────────────────────
        story.append(Spacer(1, 0.3*cm))
        story.append(Paragraph("Executive Summary", h2))
        story.append(Paragraph(ai.get("executive_summary","N/A").replace("\n","<br/>"), body))
        story.append(Spacer(1, 0.3*cm))

        # ── Threat Severity Table ──────────────────────────
        story.append(Paragraph("Threat Severity Breakdown", h2))
        sev_data = [
            ["Severity", "Count", "Status"],
            ["🔴 CRITICAL", str(sev.get("CRITICAL",0)), "IMMEDIATE ACTION"],
            ["🟠 HIGH",     str(sev.get("HIGH",0)),     "URGENT"],
            ["🟡 MEDIUM",   str(sev.get("MEDIUM",0)),   "INVESTIGATE"],
            ["🔵 LOW",      str(sev.get("LOW",0)),       "MONITOR"],
            ["TOTAL",       str(anal.get("total_threats",0)), ""],
        ]
        sev_table = Table(sev_data, colWidths=[5*cm, 4*cm, 8*cm])
        sev_table.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0),  HexColor("#1e293b")),
            ("TEXTCOLOR",   (0,0),(-1,0),  white),
            ("FONTNAME",    (0,0),(-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),["#ffffff","#f8fafc"]),
            ("GRID",        (0,0),(-1,-1), 0.5, HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 6),
            ("FONTNAME",    (0,-1),(-1,-1),"Helvetica-Bold"),
        ]))
        story.append(sev_table)
        story.append(Spacer(1, 0.3*cm))

        # ── Incident Summary ──────────────────────────────
        story.append(Paragraph("Incident Summary", h2))
        inc_data = [
            ["Metric",         "Value"],
            ["Total Incidents", str(anal.get("total_incidents",0))],
            ["Open",           str(anal.get("open_incidents",0))],
            ["Resolved",       str(anal.get("resolved_incidents",0))],
            ["Avg Risk Score", f"{anal.get('average_risk_score',0)}/100"],
        ]
        inc_table = Table(inc_data, colWidths=[8*cm, 9*cm])
        inc_table.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0),  HexColor("#1e293b")),
            ("TEXTCOLOR",   (0,0),(-1,0),  white),
            ("FONTNAME",    (0,0),(-1,0),  "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),["#ffffff","#f8fafc"]),
            ("GRID",        (0,0),(-1,-1), 0.5, HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 6),
        ]))
        story.append(inc_table)
        story.append(Spacer(1, 0.3*cm))

        # ── Top Users + IPs ───────────────────────────────
        story.append(Paragraph("Top Threat Actors", h2))
        user_data = [["Username", "Event Count"]]
        for u in anal.get("top_users", [])[:5]:
            user_data.append([u["name"], str(u["count"])])
        if len(user_data) > 1:
            u_table = Table(user_data, colWidths=[10*cm, 7*cm])
            u_table.setStyle(TableStyle([
                ("BACKGROUND", (0,0),(-1,0), HexColor("#1e293b")),
                ("TEXTCOLOR",  (0,0),(-1,0), white),
                ("FONTNAME",   (0,0),(-1,0), "Helvetica-Bold"),
                ("FONTSIZE",   (0,0),(-1,-1), 9),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),["#ffffff","#f8fafc"]),
                ("GRID",       (0,0),(-1,-1), 0.5, HexColor("#e2e8f0")),
                ("PADDING",    (0,0),(-1,-1), 6),
            ]))
            story.append(u_table)
        story.append(Spacer(1, 0.3*cm))

        # ── Critical Findings ─────────────────────────────
        findings = ai.get("critical_findings", [])
        if findings:
            story.append(Paragraph("Critical Findings", h2))
            for cf in findings[:5]:
                story.append(Paragraph(f"⚠ {cf}", body))
            story.append(Spacer(1, 0.3*cm))

        # ── Recommendations ───────────────────────────────
        story.append(Paragraph("Security Recommendations", h2))
        for i, rec in enumerate(ai.get("recommendations", [])[:10], 1):
            story.append(Paragraph(f"{i}. {rec}", body))
        story.append(Spacer(1, 0.3*cm))

        # ── Recent Threats Table ──────────────────────────
        story.append(PageBreak())
        story.append(Paragraph("Recent Threats (Top 20)", h2))
        th_data = [["Threat ID", "Event Type", "Severity", "User", "Source IP", "Time"]]
        for t in data.get("threats", [])[:20]:
            th_data.append([
                t.get("threat_id","")[:12],
                t.get("event_type","")[:20],
                t.get("severity",""),
                t.get("username","")[:12],
                t.get("source_ip","")[:15],
                t.get("timestamp","")[:16],
            ])
        if len(th_data) > 1:
            th_table = Table(th_data, colWidths=[2.5*cm,4*cm,2*cm,2.5*cm,3*cm,3*cm])
            th_table.setStyle(TableStyle([
                ("BACKGROUND",  (0,0),(-1,0),  HexColor("#1e293b")),
                ("TEXTCOLOR",   (0,0),(-1,0),  white),
                ("FONTNAME",    (0,0),(-1,0),  "Helvetica-Bold"),
                ("FONTSIZE",    (0,0),(-1,-1), 7),
                ("ROWBACKGROUNDS",(0,1),(-1,-1),["#ffffff","#f8fafc"]),
                ("GRID",        (0,0),(-1,-1), 0.3, HexColor("#e2e8f0")),
                ("PADDING",     (0,0),(-1,-1), 4),
            ]))
            story.append(th_table)

        # ── Footer ────────────────────────────────────────
        story.append(Spacer(1, cm))
        story.append(HRFlowable(width="100%", color=HexColor("#e2e8f0"), thickness=0.5))
        story.append(Paragraph("AWS Cloud SOC Platform — Confidential Security Report", center))
        story.append(Paragraph(f"Generated: {data.get('generated_at','')}", center))

        doc.build(story)
        return buffer.getvalue()

    except ImportError:
        # Fallback: if reportlab not installed, return txt as bytes
        logger.warning("reportlab not installed — returning TXT fallback")
        txt = _generate_txt(data, report_type, generated_by)
        return txt.encode("utf-8")


def _generate_docx(data: dict, report_type: str, generated_by: str) -> bytes:
    """Generate DOCX using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        ai  = data.get("ai_summary", {})
        anal = data.get("analytics", {})
        sev  = anal.get("severity_breakdown", {})

        def add_heading(text, level=1, color=(0,100,200)):
            h = doc.add_heading(text, level=level)
            h.runs[0].font.color.rgb = RGBColor(*color)
            return h

        def add_paragraph(text, bold=False, size=10):
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = bold
            run.font.size = Pt(size)
            return p

        # Cover
        title = doc.add_heading("AWS CLOUD SECURITY OPERATIONS CENTER", 0)
        title.runs[0].font.color.rgb = RGBColor(76, 201, 240)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        sub = doc.add_heading(f"{report_type.upper().replace('_',' ')} REPORT", 1)
        sub.runs[0].font.color.rgb = RGBColor(76, 201, 240)
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Metadata table
        meta_table = doc.add_table(rows=6, cols=2)
        meta_table.style = "Table Grid"
        rows = meta_table.rows
        meta_data = [
            ("Generated By",  generated_by),
            ("Generated At",  data.get("generated_at","")),
            ("Security Score", f"{ai.get('security_score','N/A')}/100"),
            ("Risk Level",    ai.get("risk_level","N/A")),
            ("Total Threats", str(anal.get("total_threats",0))),
            ("Open Incidents",str(anal.get("open_incidents",0))),
        ]
        for i, (k, v) in enumerate(meta_data):
            rows[i].cells[0].text = k
            rows[i].cells[1].text = v
            rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Executive Summary
        add_heading("Executive Summary", 1)
        add_paragraph(ai.get("executive_summary","N/A"))
        doc.add_paragraph()

        # Threat Summary
        add_heading("Threat Severity Breakdown", 2)
        sev_table = doc.add_table(rows=6, cols=2)
        sev_table.style = "Table Grid"
        sev_rows = [
            ("CRITICAL", str(sev.get("CRITICAL",0))),
            ("HIGH",     str(sev.get("HIGH",0))),
            ("MEDIUM",   str(sev.get("MEDIUM",0))),
            ("LOW",      str(sev.get("LOW",0))),
            ("TOTAL",    str(anal.get("total_threats",0))),
        ]
        sev_table.rows[0].cells[0].text = "Severity"
        sev_table.rows[0].cells[1].text = "Count"
        for i, (k, v) in enumerate(sev_rows, 1):
            sev_table.rows[i].cells[0].text = k
            sev_table.rows[i].cells[1].text = v

        doc.add_paragraph()

        # Incident Summary
        add_heading("Incident Summary", 2)
        inc_data = [
            ("Total", str(anal.get("total_incidents",0))),
            ("Open",  str(anal.get("open_incidents",0))),
            ("Resolved", str(anal.get("resolved_incidents",0))),
        ]
        inc_table = doc.add_table(rows=len(inc_data)+1, cols=2)
        inc_table.style = "Table Grid"
        inc_table.rows[0].cells[0].text = "Status"
        inc_table.rows[0].cells[1].text = "Count"
        for i, (k, v) in enumerate(inc_data, 1):
            inc_table.rows[i].cells[0].text = k
            inc_table.rows[i].cells[1].text = v

        doc.add_paragraph()

        # Recommendations
        add_heading("Security Recommendations", 2)
        for i, rec in enumerate(ai.get("recommendations",[])[:10], 1):
            add_paragraph(f"{i}. {rec}")

        doc.add_paragraph()

        # Recent Threats
        add_heading("Recent Threats (Top 20)", 2)
        threats = data.get("threats", [])[:20]
        if threats:
            t_table = doc.add_table(rows=len(threats)+1, cols=5)
            t_table.style = "Table Grid"
            headers = ["Threat ID","Event Type","Severity","Username","Source IP"]
            for j, h in enumerate(headers):
                t_table.rows[0].cells[j].text = h
                t_table.rows[0].cells[j].paragraphs[0].runs[0].bold = True
            for i, t in enumerate(threats, 1):
                t_table.rows[i].cells[0].text = t.get("threat_id","")[:12]
                t_table.rows[i].cells[1].text = t.get("event_type","")[:25]
                t_table.rows[i].cells[2].text = t.get("severity","")
                t_table.rows[i].cells[3].text = t.get("username","")[:15]
                t_table.rows[i].cells[4].text = t.get("source_ip","")

        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph("AWS Cloud SOC Platform — Confidential Security Report")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.runs[0].font.color.rgb = RGBColor(107, 114, 128)
        footer.runs[0].font.size = Pt(8)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    except ImportError:
        logger.warning("python-docx not installed — returning TXT fallback")
        txt = _generate_txt(data, report_type, generated_by)
        return txt.encode("utf-8")